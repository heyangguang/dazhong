#!/usr/bin/env python3
from docx import Document

# 读取生成的Word文档
doc = Document('test_report_fixed.docx')

print("=== Word文档内容检查 ===\n")

# 检查段落内容
for i, paragraph in enumerate(doc.paragraphs[:10]):  # 检查前10个段落
    if paragraph.text.strip():
        print(f"段落 {i+1}: {paragraph.text[:50]}...")
        # 检查是否有乱码字符
        if any(ord(char) == 0xFFFD or char == '□' for char in paragraph.text):
            print("  ⚠️  发现可能的乱码字符!")
        else:
            print("  ✓ 文本正常")

print("\n=== 字体信息 ===")
# 检查字体设置
for paragraph in doc.paragraphs[:5]:
    if paragraph.text.strip():
        for run in paragraph.runs:
            if run.text.strip():
                font_name = run.font.name if run.font.name else "默认"
                print(f"文本: '{run.text[:20]}...' - 字体: {font_name}")
                break