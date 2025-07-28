from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
from datetime import datetime
from models import db, User, Vehicle, Activity, Evaluator, Evaluation, Category, ActivityEvaluator
from docx import Document
from docx.shared import Inches
import io
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///evaluation.db'
app.config['UPLOAD_FOLDER'] = 'static/uploads'

db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'admin_login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/')
def index():
    activities = Activity.query.all()
    return render_template('index.html', activities=activities)

@app.route('/categories/<int:activity_id>/<int:vehicle_id>')
def categories(activity_id, vehicle_id):
    activity = Activity.query.get_or_404(activity_id)
    vehicle = Vehicle.query.get_or_404(vehicle_id)
    categories = Category.query.filter_by(is_active=True).order_by(Category.order).all()
    return render_template('categories.html', activity=activity, vehicle=vehicle, categories=categories)

@app.route('/evaluation/<int:activity_id>/<int:vehicle_id>/<category>')
def evaluation(activity_id, vehicle_id, category):
    activity = Activity.query.get_or_404(activity_id)
    vehicle = Vehicle.query.get_or_404(vehicle_id)
    
    # 只获取关联到该活动的评价人
    activity_evaluators = ActivityEvaluator.query.filter_by(activity_id=activity_id).all()
    evaluator_ids = [ae.evaluator_id for ae in activity_evaluators]
    
    if evaluator_ids:
        evaluators = Evaluator.query.filter(Evaluator.id.in_(evaluator_ids)).all()
    else:
        # 如果活动没有关联评价人，则显示所有评价人
        evaluators = Evaluator.query.all()
    
    return render_template('evaluation.html', 
                         activity=activity, 
                         vehicle=vehicle, 
                         category=category,
                         evaluators=evaluators)

@app.route('/api/activities')
def api_activities():
    activities = Activity.query.all()
    return jsonify([{
        'id': a.id,
        'name': a.name,
        'vehicle_id': a.vehicle_id,
        'vehicle_name': a.vehicle.name if a.vehicle else ''
    } for a in activities])

@app.route('/api/save_evaluation', methods=['POST'])
def save_evaluation():
    data = request.json
    evaluation = Evaluation(
        activity_id=data['activity_id'],
        vehicle_id=data['vehicle_id'],
        evaluator_id=data['evaluator_id'],
        category=data['category'],
        score=data['score'],
        content=data['content']
    )
    db.session.add(evaluation)
    db.session.commit()
    return jsonify({'status': 'success', 'id': evaluation.id})

@app.route('/api/evaluations/<int:activity_id>/<int:vehicle_id>')
def get_evaluations(activity_id, vehicle_id):
    evaluations = Evaluation.query.filter_by(
        activity_id=activity_id,
        vehicle_id=vehicle_id
    ).order_by(Evaluation.created_at.desc()).all()
    
    return jsonify([{
        'id': e.id,
        'category': e.category,
        'score': e.score,
        'content': e.content,
        'evaluator_name': e.evaluator.name,
        'created_at': e.created_at.strftime('%Y-%m-%d %H:%M')
    } for e in evaluations])

@app.route('/api/evaluation_counts/<int:activity_id>/<int:vehicle_id>')
def get_evaluation_counts(activity_id, vehicle_id):
    counts = {}
    categories = Category.query.filter_by(is_active=True).all()
    
    for category in categories:
        count = Evaluation.query.filter_by(
            activity_id=activity_id,
            vehicle_id=vehicle_id,
            category=category.name
        ).count()
        counts[category.name] = count
    
    return jsonify(counts)

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('admin_dashboard'))
    return render_template('admin/login.html')

@app.route('/admin/logout')
@login_required
def admin_logout():
    logout_user()
    return redirect(url_for('admin_login'))

@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    vehicles = Vehicle.query.all()
    activities = Activity.query.all()
    evaluators = Evaluator.query.all()
    categories = Category.query.order_by(Category.order).all()
    return render_template('admin/dashboard.html',
                         vehicles=vehicles,
                         activities=activities,
                         evaluators=evaluators,
                         categories=categories)

@app.route('/admin/evaluations/<int:activity_id>')
@login_required
def view_evaluations(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    evaluations = Evaluation.query.filter_by(activity_id=activity_id).order_by(Evaluation.created_at.desc()).all()
    categories = Category.query.filter_by(is_active=True).order_by(Category.order).all()
    
    # 统计数据
    stats = {
        'total': len(evaluations),
        'by_category': {},
        'avg_score': 0
    }
    
    for category in categories:
        cat_evals = [e for e in evaluations if e.category == category.name]
        stats['by_category'][category.name] = {
            'count': len(cat_evals),
            'avg_score': sum(e.score for e in cat_evals) / len(cat_evals) if cat_evals else 0
        }
    
    stats['avg_score'] = sum(e.score for e in evaluations) / len(evaluations) if evaluations else 0
    
    return render_template('admin/evaluations.html',
                         activity=activity,
                         evaluations=evaluations,
                         categories=categories,
                         stats=stats)

@app.route('/admin/vehicle/add', methods=['POST'])
@login_required
def add_vehicle():
    name = request.form.get('name')
    model = request.form.get('model')
    
    vehicle = Vehicle(name=name, model=model)
    
    if 'info_file' in request.files:
        file = request.files['info_file']
        if file and file.filename.endswith('.pdf'):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            vehicle.info_file = filename
    
    db.session.add(vehicle)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/activity/add', methods=['POST'])
@login_required
def add_activity():
    name = request.form.get('name')
    date = datetime.strptime(request.form.get('date'), '%Y-%m-%d').date()
    vehicle_id = request.form.get('vehicle_id')
    evaluator_ids = request.form.getlist('evaluator_ids')
    
    activity = Activity(name=name, date=date, vehicle_id=vehicle_id)
    db.session.add(activity)
    db.session.commit()
    
    # 关联评价人
    for evaluator_id in evaluator_ids:
        if evaluator_id:
            ae = ActivityEvaluator(activity_id=activity.id, evaluator_id=int(evaluator_id))
            db.session.add(ae)
    db.session.commit()
    
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/evaluator/add', methods=['POST'])
@login_required
def add_evaluator():
    name = request.form.get('name')
    department = request.form.get('department')
    
    evaluator = Evaluator(name=name, department=department)
    db.session.add(evaluator)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/vehicle/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_vehicle(id):
    vehicle = Vehicle.query.get_or_404(id)
    
    if request.method == 'POST':
        vehicle.name = request.form.get('name')
        vehicle.model = request.form.get('model')
        
        if 'info_file' in request.files:
            file = request.files['info_file']
            if file and file.filename.endswith('.pdf'):
                # 删除旧文件
                if vehicle.info_file and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], vehicle.info_file)):
                    os.remove(os.path.join(app.config['UPLOAD_FOLDER'], vehicle.info_file))
                
                # 保存新文件
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                vehicle.info_file = filename
        
        db.session.commit()
        return redirect(url_for('admin_dashboard'))
    
    return jsonify({
        'id': vehicle.id,
        'name': vehicle.name,
        'model': vehicle.model,
        'info_file': vehicle.info_file
    })

@app.route('/admin/vehicle/delete/<int:id>')
@login_required
def delete_vehicle(id):
    vehicle = Vehicle.query.get_or_404(id)
    
    # 删除关联的PDF文件
    if vehicle.info_file and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], vehicle.info_file)):
        os.remove(os.path.join(app.config['UPLOAD_FOLDER'], vehicle.info_file))
    
    db.session.delete(vehicle)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/evaluator/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_evaluator(id):
    evaluator = Evaluator.query.get_or_404(id)
    
    if request.method == 'POST':
        evaluator.name = request.form.get('name')
        evaluator.department = request.form.get('department')
        db.session.commit()
        return redirect(url_for('admin_dashboard'))
    
    return jsonify({
        'id': evaluator.id,
        'name': evaluator.name,
        'department': evaluator.department
    })

@app.route('/admin/evaluator/delete/<int:id>')
@login_required
def delete_evaluator(id):
    evaluator = Evaluator.query.get_or_404(id)
    db.session.delete(evaluator)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/activity/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_activity(id):
    activity = Activity.query.get_or_404(id)
    
    if request.method == 'POST':
        activity.name = request.form.get('name')
        activity.date = datetime.strptime(request.form.get('date'), '%Y-%m-%d').date()
        activity.vehicle_id = request.form.get('vehicle_id')
        
        # 更新关联的评价人
        evaluator_ids = request.form.getlist('evaluator_ids')
        
        # 删除现有关联
        ActivityEvaluator.query.filter_by(activity_id=activity.id).delete()
        
        # 添加新关联
        for evaluator_id in evaluator_ids:
            if evaluator_id:
                ae = ActivityEvaluator(activity_id=activity.id, evaluator_id=int(evaluator_id))
                db.session.add(ae)
        
        db.session.commit()
        return redirect(url_for('admin_dashboard'))
    
    # 获取活动的评价人ID列表
    evaluator_ids = [ae.evaluator_id for ae in activity.activity_evaluators]
    
    return jsonify({
        'id': activity.id,
        'name': activity.name,
        'date': activity.date.strftime('%Y-%m-%d'),
        'vehicle_id': activity.vehicle_id,
        'evaluator_ids': evaluator_ids
    })

@app.route('/admin/activity/delete/<int:id>')
@login_required
def delete_activity(id):
    activity = Activity.query.get_or_404(id)
    
    # 删除关联的评价人
    ActivityEvaluator.query.filter_by(activity_id=activity.id).delete()
    
    db.session.delete(activity)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/category/add', methods=['POST'])
@login_required
def add_category():
    name = request.form.get('name')
    name_en = request.form.get('name_en')
    icon = request.form.get('icon', 'default')
    order = request.form.get('order', 0, type=int)
    
    category = Category(name=name, name_en=name_en, icon=icon, order=order)
    db.session.add(category)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/category/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_category(id):
    category = Category.query.get_or_404(id)
    
    if request.method == 'POST':
        category.name = request.form.get('name')
        category.name_en = request.form.get('name_en')
        category.icon = request.form.get('icon', 'default')
        category.order = request.form.get('order', 0, type=int)
        category.is_active = request.form.get('is_active') == 'on'
        db.session.commit()
        return redirect(url_for('admin_dashboard'))
    
    # GET 请求返回分类数据
    return jsonify({
        'id': category.id,
        'name': category.name,
        'name_en': category.name_en,
        'icon': category.icon,
        'order': category.order,
        'is_active': category.is_active
    })

@app.route('/admin/category/delete/<int:id>')
@login_required
def delete_category(id):
    category = Category.query.get_or_404(id)
    # 检查是否有评价使用此分类
    if category.evaluations.count() > 0:
        return jsonify({'error': '该分类下有评价记录，无法删除'}), 400
    db.session.delete(category)
    db.session.commit()
    return redirect(url_for('admin_dashboard'))

@app.route('/pdf_viewer/<int:vehicle_id>')
def pdf_viewer(vehicle_id):
    vehicle = Vehicle.query.get_or_404(vehicle_id)
    return render_template('pdf_viewer.html', vehicle=vehicle)

@app.route('/vehicle/<int:vehicle_id>/pdf')
def vehicle_pdf(vehicle_id):
    vehicle = Vehicle.query.get_or_404(vehicle_id)
    if vehicle.info_file:
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], vehicle.info_file)
        if os.path.exists(pdf_path):
            return send_file(pdf_path, mimetype='application/pdf')
    
    # 如果没有上传的PDF，返回404
    return "PDF文件未找到", 404

@app.route('/admin/export/<int:activity_id>')
@login_required
def export_report(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    evaluations = Evaluation.query.filter_by(activity_id=activity_id).all()
    
    doc = Document()
    
    # 设置文档的默认字体为支持中文的字体
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    
    # 设置文档样式 - 使用更兼容的方式
    # 创建一个辅助函数来安全地设置字体
    def set_font_safe(run, font_name='Arial', size=11):
        """安全地设置字体，避免因字体不存在而出错"""
        try:
            run.font.name = font_name
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            if size:
                run.font.size = Pt(size)
        except:
            # 如果设置失败，使用默认配置
            pass
    
    # 设置默认样式
    try:
        # 设置Normal样式
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    except:
        pass
    
    # 添加标题
    title = doc.add_heading(f'{activity.name} - 评价报告', 0)
    title.alignment = 1  # 居中对齐
    # 设置标题字体
    for run in title.runs:
        set_font_safe(run, 'Calibri', 16)
    
    # 添加基本信息
    info_para = doc.add_paragraph()
    info_para.add_run(f'活动日期: {activity.date}\n').bold = True
    info_para.add_run(f'车型: {activity.vehicle.name}\n').bold = True
    info_para.add_run(f'报告生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # 设置段落字体
    for run in info_para.runs:
        set_font_safe(run, 'Calibri', 12)
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加评价内容
    categories = ['动力总成', '底盘', '内饰', '外饰', '声学', '电子电器', '其它']
    for category in categories:
        cat_evals = [e for e in evaluations if e.category == category]
        if cat_evals:
            # 分类标题
            cat_heading = doc.add_heading(category, 1)
            for run in cat_heading.runs:
                set_font_safe(run, 'Calibri', 14)
            
            # 评价内容
            for eval in cat_evals:
                # 评价人和评分
                eval_info = doc.add_paragraph()
                eval_info.add_run(f'评价人: ').bold = True
                eval_info.add_run(f'{eval.evaluator.name}')
                eval_info.add_run(f'    评分: ').bold = True
                eval_info.add_run(f'{eval.score}/10')
                
                # 设置字体和颜色
                for run in eval_info.runs:
                    set_font_safe(run, 'Calibri', 11)
                    
                    # 根据评分设置颜色
                    if '评分:' in run.text and eval.score:
                        if eval.score <= 3:
                            run.font.color.rgb = RGBColor(220, 53, 69)  # 红色
                        elif eval.score <= 5:
                            run.font.color.rgb = RGBColor(253, 126, 20)  # 橙色
                        elif eval.score <= 7:
                            run.font.color.rgb = RGBColor(255, 193, 7)  # 黄色
                        else:
                            run.font.color.rgb = RGBColor(40, 167, 69)  # 绿色
                
                # 评价内容
                content_para = doc.add_paragraph()
                content_para.add_run('评价内容: ').bold = True
                # 清理内容中的HTML标签
                import re
                clean_content = re.sub('<.*?>', '', eval.content)
                content_para.add_run(clean_content)
                
                # 设置字体
                for run in content_para.runs:
                    set_font_safe(run, 'Calibri', 11)
                
                # 添加时间戳
                time_para = doc.add_paragraph()
                time_para.add_run(f'评价时间: {eval.created_at.strftime("%Y-%m-%d %H:%M:%S")}')
                for run in time_para.runs:
                    set_font_safe(run, 'Calibri', 10)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                
                # 添加分隔线
                doc.add_paragraph('─' * 50)
    
    # 添加总结信息
    doc.add_page_break()
    summary_heading = doc.add_heading('评价总结', 1)
    for run in summary_heading.runs:
        set_font_safe(run, 'Calibri', 14)
    
    # 统计各分类评价数量
    summary_para = doc.add_paragraph()
    summary_para.add_run('各分类评价统计:\n').bold = True
    for category in categories:
        cat_evals = [e for e in evaluations if e.category == category]
        if cat_evals:
            avg_score = sum(e.score for e in cat_evals) / len(cat_evals)
            summary_para.add_run(f'{category}: {len(cat_evals)}条评价，平均分: {avg_score:.1f}\n')
    
    for run in summary_para.runs:
        set_font_safe(run, 'Calibri', 11)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'{activity.name}_评价报告_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # 初始化管理员账号
        if not User.query.filter_by(username='admin').first():
            admin = User(
                username='admin',
                password=generate_password_hash('admin123')
            )
            db.session.add(admin)
            db.session.commit()
        
        # 初始化车型数据
        if Vehicle.query.count() == 0:
            # 创建示例PDF
            os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
            sample_pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], 'sample_vehicle_info.pdf')
            
            if not os.path.exists(sample_pdf_path):
                c = canvas.Canvas(sample_pdf_path, pagesize=A4)
                width, height = A4
                
                c.setFont("Helvetica-Bold", 24)
                c.drawString(50, height - 100, "Vehicle Technical Specifications")
                
                c.setFont("Helvetica", 14)
                y_position = height - 150
                
                info_lines = [
                    "Model: Pos.6-vw316/8 CM_B A SUVe Black Label",
                    "Type: Luxury SUV",
                    "Engine: 2.0L Turbocharged",
                    "Power Output: 252 HP @ 5,200 rpm",
                    "Torque: 370 Nm @ 1,600-4,500 rpm",
                    "Drivetrain: All-Wheel Drive (AWD)",
                    "Transmission: 8-Speed Automatic",
                    "Fuel Type: Premium Gasoline",
                    "Seating Capacity: 7 passengers",
                    "Overall Length: 4,925 mm",
                    "Overall Width: 1,985 mm",
                    "Overall Height: 1,770 mm",
                    "Wheelbase: 2,915 mm",
                    "Curb Weight: 2,100 kg",
                    "Fuel Tank Capacity: 80 liters",
                    "0-100 km/h Acceleration: 7.5 seconds",
                    "Top Speed: 210 km/h",
                    "Fuel Economy (Combined): 8.5 L/100km"
                ]
                
                for line in info_lines:
                    c.drawString(50, y_position, line)
                    y_position -= 25
                
                c.save()
            
            vehicles = [
                Vehicle(name='Pos.6-vw316/8 CM_B A SUVe Black Label', model='SUVe Black Label', info_file='sample_vehicle_info.pdf'),
                Vehicle(name='BMW X5 xDrive40i', model='X5'),
                Vehicle(name='Mercedes-Benz GLE 450', model='GLE'),
                Vehicle(name='Audi Q7 55 TFSI', model='Q7')
            ]
            for v in vehicles:
                db.session.add(v)
            db.session.commit()
        
        # 初始化分类数据
        if Category.query.count() == 0:
            default_categories = [
                {'name': '动力总成', 'name_en': 'Powertrain', 'icon': 'powertrain', 'order': 1},
                {'name': '底盘', 'name_en': 'Chassis', 'icon': 'chassis', 'order': 2},
                {'name': '内饰', 'name_en': 'Interior', 'icon': 'interior', 'order': 3},
                {'name': '外饰', 'name_en': 'Exterior', 'icon': 'exterior', 'order': 4},
                {'name': '声学', 'name_en': 'Acoustic', 'icon': 'acoustic', 'order': 5},
                {'name': '电子电器', 'name_en': 'Electronic', 'icon': 'electronic', 'order': 6},
                {'name': '其它', 'name_en': 'Others', 'icon': 'other', 'order': 7}
            ]
            for cat_data in default_categories:
                category = Category(**cat_data)
                db.session.add(category)
            db.session.commit()
            print("分类数据初始化完成！")
        
        # 初始化评价人数据
        if Evaluator.query.count() == 0:
            evaluators = [
                Evaluator(name='张三', department='动力总成部'),
                Evaluator(name='李四', department='底盘部'),
                Evaluator(name='王五', department='内饰部'),
                Evaluator(name='赵六', department='外饰部'),
                Evaluator(name='陈七', department='电子电器部'),
                Evaluator(name='刘八', department='质量部'),
                Evaluator(name='孙九', department='研发部'),
                Evaluator(name='周十', department='测试部')
            ]
            for e in evaluators:
                db.session.add(e)
            db.session.commit()
        
        # 初始化活动数据
        if Activity.query.count() == 0:
            activities = [
                Activity(
                    name='CFF Warm主试车 - 2024年1月',
                    date=datetime(2024, 1, 15).date(),
                    vehicle_id=1
                ),
                Activity(
                    name='冬季测试活动 - 2024年2月',
                    date=datetime(2024, 2, 20).date(),
                    vehicle_id=2
                ),
                Activity(
                    name='春季评测活动 - 2024年3月',
                    date=datetime(2024, 3, 10).date(),
                    vehicle_id=3
                )
            ]
            for a in activities:
                db.session.add(a)
            db.session.commit()
        
        # 初始化一些示例评价数据
        if Evaluation.query.count() == 0:
            sample_evaluations = [
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=1,
                    category='动力总成',
                    score=8,
                    content='动力输出平顺，加速响应良好，但在高速行驶时噪音略大。建议优化发动机隔音材料。'
                ),
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=2,
                    category='底盘',
                    score=9,
                    content='悬架调校偏向舒适，过弯侧倾控制良好。刹车脚感线性，制动力充足。'
                ),
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=3,
                    category='内饰',
                    score=7,
                    content='内饰用料考究，做工精细。中控屏幕反应灵敏，但部分按键位置不够合理。'
                ),
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=5,
                    category='电子电器',
                    score=8,
                    content='车机系统流畅，功能丰富。自动驾驶辅助功能表现出色，但语音识别准确率有待提高。'
                ),
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=5,
                    category='电子电器',
                    score=7,
                    content='空调制冷效果良好，分区控制精准。座椅加热功能响应迅速。'
                ),
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=4,
                    category='声学',
                    score=6,
                    content='整体隔音效果一般，高速风噪明显。建议加强A柱和后视镜区域的密封。'
                ),
                Evaluation(
                    activity_id=1,
                    vehicle_id=1,
                    evaluator_id=6,
                    category='声学',
                    score=7,
                    content='音响系统音质出色，低音下潜有力。但在颠簸路面会有轻微异响。'
                )
            ]
            for eval in sample_evaluations:
                db.session.add(eval)
            db.session.commit()
            
        print("数据初始化完成！")
        print("默认管理员账号: admin / admin123")
        print("访问地址: http://localhost:5000")
    
    app.run(debug=True, host='0.0.0.0', port=5000)